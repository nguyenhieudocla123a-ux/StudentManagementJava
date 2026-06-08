from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()
section = doc.sections[0]
section.left_margin = Cm(2.5)
section.right_margin = Cm(2)
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2.5)

def h1(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(16)
    run.bold = True
    run.font.color.rgb = RGBColor(0, 70, 127)
    return p

def h2(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(13)
    run.bold = True
    run.font.color.rgb = RGBColor(0, 112, 192)
    return p

def h3(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(11)
    run.bold = True
    run.font.color.rgb = RGBColor(68, 114, 196)
    return p

def normal(text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)
    run.bold = bold
    return p

def code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0, 0, 0)
    # background gray
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F2F2F2')
    pPr.append(shd)
    return p

def diagram(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.5)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0, 100, 0)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'E8F5E9')
    pPr.append(shd)
    return p

def sep():
    doc.add_paragraph('─' * 80)

# ============================================================
# TITLE
# ============================================================
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run('JAVA CORE → SPRING BOOT\nHỌC QUA DỰ ÁN THỰC TẾ')
r.font.name = 'Consolas'
r.font.size = Pt(18)
r.bold = True
r.font.color.rgb = RGBColor(0, 70, 127)
doc.add_paragraph()

# ============================================================
# PHẦN 1: JAVA CORE
# ============================================================
h1('PHẦN 1: JAVA CORE')
sep()

# 1.1 Biến và kiểu dữ liệu
h2('1.1 Biến và Kiểu Dữ Liệu')
diagram('''
┌─────────────────────────────────────────────────────┐
│                  KIỂU DỮ LIỆU JAVA                  │
├──────────────────────┬──────────────────────────────┤
│  Primitive (nguyên)  │  Reference (tham chiếu)      │
├──────────────────────┼──────────────────────────────┤
│  int    → 1, 2, 3    │  String  → "Hello"           │
│  double → 3.14       │  int[]   → {1,2,3}           │
│  boolean→ true/false │  Object  → new GiangVien()   │
│  char   → 'A'        │  List    → ArrayList<>()     │
└──────────────────────┴──────────────────────────────┘
''')
code_block('''// Khai báo biến
int soTinChi = 3;
double diemTongKet = 8.5;
boolean isPass = true;
String maGV = "GV001";

// Từ dự án thực tế
String maSV = "SV001";
float diemQT = 8.0f;
float diemGK = 7.5f;
float diemCK = 9.0f;
float diemTK = diemQT * 0.1f + diemGK * 0.3f + diemCK * 0.6f; // = 8.45
''')

# 1.2 Class và Object
h2('1.2 Class và Object')
diagram('''
CLASS = Bản thiết kế          OBJECT = Sản phẩm thực tế
┌─────────────────────┐       ┌─────────────────────┐
│   class GiangVien   │  →→→  │  gv = new GiangVien │
│  ─────────────────  │       │  ─────────────────  │
│  String maGV        │       │  maGV = "GV001"     │
│  String hoTen       │       │  hoTen = "Nguyen A" │
│  String email       │       │  email = "a@edu.vn" │
│  ─────────────────  │       │  ─────────────────  │
│  getMaGV()          │       │  gv.getMaGV()       │
│  setMaGV()          │       │  gv.setMaGV("GV002")│
└─────────────────────┘       └─────────────────────┘
''')
code_block('''// Định nghĩa class
public class GiangVien {
    private String maGV;
    private String hoTen;

    // Constructor
    public GiangVien(String maGV, String hoTen) {
        this.maGV = maGV;
        this.hoTen = hoTen;
    }

    // Getter / Setter
    public String getMaGV() { return maGV; }
    public void setMaGV(String maGV) { this.maGV = maGV; }
}

// Tạo object
GiangVien gv = new GiangVien("GV001", "Nguyen Van A");
System.out.println(gv.getMaGV()); // GV001
''')

# 1.3 Interface
h2('1.3 Interface')
diagram('''
Interface = Hợp đồng (quy định phải làm gì)
Class = Người thực hiện hợp đồng

┌──────────────────────┐
│  <<interface>>       │
│  Printable           │
│  ─────────────────   │
│  + print()  [trống]  │
└──────────┬───────────┘
           │ implements
    ┌──────┴──────┐
    │             │
┌───▼───┐   ┌────▼────┐
│Report │   │Invoice  │
│print()│   │print()  │
└───────┘   └─────────┘
''')
code_block('''// Interface
public interface Printable {
    void print(); // chỉ khai báo, không có code
}

// Class implement interface
public class Report implements Printable {
    @Override
    public void print() {
        System.out.println("In báo cáo...");
    }
}
''')

# 1.4 Collections
h2('1.4 Collections (Danh sách, Map)')
diagram('''
┌─────────────────────────────────────────────────────────┐
│                    JAVA COLLECTIONS                     │
├─────────────────┬───────────────────┬───────────────────┤
│   ArrayList     │      HashMap      │      HashSet      │
├─────────────────┼───────────────────┼───────────────────┤
│ [GV001, GV002,  │ "GV001" → GiangV  │ {GV001, GV002,   │
│  GV003, ...]    │ "GV002" → GiangV  │  GV003} (no dup) │
│                 │ "GV003" → GiangV  │                   │
│ Có thứ tự       │ Key-Value pair    │ Không trùng lặp   │
│ Cho phép trùng  │ Tìm nhanh theo key│                   │
└─────────────────┴───────────────────┴───────────────────┘
''')
code_block('''// ArrayList - từ dự án thực tế
ArrayList<GiangVien> list = dao.getAllGiangVien();
for (GiangVien gv : list) {
    System.out.println(gv.getMaGV() + " - " + gv.getHoTen());
}

// HashMap
HashMap<String, GiangVien> map = new HashMap<>();
map.put("GV001", gv1);
GiangVien found = map.get("GV001"); // tìm nhanh

// Stream API (Java 8+)
boolean hasGV001 = list.stream()
    .anyMatch(gv -> "GV001".equals(gv.getMaGV()));
''')

# 1.5 Exception
h2('1.5 Exception Handling')
diagram('''
LUỒNG XỬ LÝ EXCEPTION:

try {                          ┌─────────────────┐
    // code có thể lỗi    →→→  │  Nếu lỗi xảy ra │
    dao.themSinhVien(sv);      └────────┬────────┘
}                                       ↓
catch (RuntimeException e) {   ┌─────────────────┐
    // xử lý lỗi           ←←← │  Bắt lỗi ở đây  │
    showError(e.getMessage());  └─────────────────┘
}
finally {                      ┌─────────────────┐
    // luôn chạy           ←←← │  Luôn thực thi  │
    closeConnection();          └─────────────────┘
}
''')
code_block('''// Từ dự án thực tế - SinhvienDao
public boolean themSinhVien(SinhVien sv) {
    try {
        String response = ApiClient.post("/students", jsonBody);
        if (JsonParser.isSuccess(response)) {
            return true;
        } else {
            String message = JsonParser.getMessage(response);
            throw new RuntimeException(message); // ném lỗi
        }
    } catch (Exception e) {
        throw new RuntimeException(e.getMessage()); // bắt và ném lại
    }
}
''')

# ============================================================
# PHẦN 2: SPRING BOOT
# ============================================================
doc.add_page_break()
h1('PHẦN 2: SPRING BOOT')
sep()

# 2.1 Kiến trúc
h2('2.1 Kiến Trúc Spring Boot (3 tầng)')
diagram('''
┌─────────────────────────────────────────────────────────────┐
│                    CLIENT (Postman / App)                    │
│              gửi HTTP Request (JSON)                        │
└──────────────────────────┬──────────────────────────────────┘
                           ↓ HTTP
┌──────────────────────────▼──────────────────────────────────┐
│                    SPRING BOOT SERVER                        │
│  ┌─────────────────────────────────────────────────────┐    │
│  │  FILTER LAYER (Security, CORS, CSRF)                │    │
│  │  Kiểm tra request hợp lệ trước khi vào controller  │    │
│  └──────────────────────┬──────────────────────────────┘    │
│                         ↓                                    │
│  ┌──────────────────────▼──────────────────────────────┐    │
│  │  CONTROLLER LAYER                                   │    │
│  │  @RestController - nhận request, trả response       │    │
│  │  AuthController, SinhVienController, ...            │    │
│  └──────────────────────┬──────────────────────────────┘    │
│                         ↓                                    │
│  ┌──────────────────────▼──────────────────────────────┐    │
│  │  REPOSITORY LAYER (JPA)                             │    │
│  │  TaiKhoanRepository, SinhVienRepository, ...        │    │
│  └──────────────────────┬──────────────────────────────┘    │
│                         ↓                                    │
└─────────────────────────┼───────────────────────────────────┘
                          ↓ SQL
┌─────────────────────────▼───────────────────────────────────┐
│                    DATABASE (SQL Server)                     │
│              TAIKHOAN, SINHVIEN, GIANGVIEN, ...             │
└─────────────────────────────────────────────────────────────┘
''')

# 2.2 Annotation
h2('2.2 Các Annotation Quan Trọng')
diagram('''
┌────────────────────────────────────────────────────────────┐
│                    ANNOTATION                              │
├──────────────────────┬─────────────────────────────────────┤
│  @RestController     │  Đây là REST API controller         │
│  @RequestMapping     │  Base URL cho controller            │
│  @GetMapping         │  Xử lý GET request                  │
│  @PostMapping        │  Xử lý POST request                 │
│  @PutMapping         │  Xử lý PUT request                  │
│  @DeleteMapping      │  Xử lý DELETE request               │
├──────────────────────┼─────────────────────────────────────┤
│  @Entity             │  Class ánh xạ với bảng DB           │
│  @Table(name="...")  │  Tên bảng trong DB                  │
│  @Id                 │  Khóa chính                         │
│  @Column(name="...") │  Tên cột trong DB                   │
├──────────────────────┼─────────────────────────────────────┤
│  @RequestBody        │  Đọc JSON từ body request           │
│  @PathVariable       │  Đọc giá trị từ URL /students/{id}  │
│  @RequestParam       │  Đọc query param ?keyword=abc       │
├──────────────────────┼─────────────────────────────────────┤
│  @CrossOrigin        │  Cho phép CORS từ domain khác       │
│  @RequiredArgsConstr │  Lombok: tạo constructor tự động    │
│  @Data               │  Lombok: tạo getter/setter tự động  │
└──────────────────────┴─────────────────────────────────────┘
''')

# 2.3 Entity
h2('2.3 Entity - Ánh xạ với Database')
diagram('''
DATABASE TABLE                    JAVA ENTITY CLASS
─────────────────                 ─────────────────────────
GIANGVIEN                         @Entity
  ma_gv    VARCHAR(10)    ←→      @Id @Column("ma_gv")
  ho_ten   NVARCHAR(100)  ←→      @Column("ho_ten")
  email    VARCHAR(100)   ←→      @Column("email")
  ma_khoa  VARCHAR(10)    ←→      @Column("ma_khoa")
''')
code_block('''// Từ dự án thực tế
@Data                          // Lombok: tự tạo getter/setter
@Entity
@Table(name = "GIANGVIEN")     // tên bảng trong DB
public class GiangVien {

    @Id
    @Column(name = "ma_gv")    // tên cột trong DB
    private String maGV;

    @Column(name = "ho_ten")
    private String hoTen;

    @Column(name = "email")
    private String email;

    @Column(name = "ma_khoa")
    private String maKhoa;

    // @Data tự tạo:
    // getMaGV(), setMaGV(), getHoTen(), setHoTen(), ...
}
''')

# 2.4 Repository
h2('2.4 Repository - Truy vấn Database')
diagram('''
JpaRepository cung cấp sẵn các method:

┌─────────────────────────────────────────────────────────┐
│  JpaRepository<GiangVien, String>                       │
├─────────────────────────────────────────────────────────┤
│  findAll()           → SELECT * FROM GIANGVIEN          │
│  findById("GV001")   → SELECT * WHERE ma_gv='GV001'     │
│  save(gv)            → INSERT hoặc UPDATE               │
│  delete(gv)          → DELETE WHERE ma_gv=?             │
│  existsById("GV001") → SELECT COUNT(*) WHERE ma_gv=?    │
│  count()             → SELECT COUNT(*) FROM GIANGVIEN   │
└─────────────────────────────────────────────────────────┘
''')
code_block('''// Từ dự án thực tế
public interface GiangVienRepository
    extends JpaRepository<GiangVien, String> {

    // JPA tự tạo SQL từ tên method:
    List<GiangVien> findByHoTenContaining(String keyword);
    // → SELECT * FROM GIANGVIEN WHERE ho_ten LIKE '%keyword%'
}

// Dùng trong Controller:
List<GiangVien> list = giangVienRepository.findAll();
Optional<GiangVien> gv = giangVienRepository.findById("GV001");
boolean exists = giangVienRepository.existsById("GV001");
''')

# 2.5 Controller
h2('2.5 Controller - Xử lý Request')
diagram('''
LUỒNG XỬ LÝ REQUEST TRONG CONTROLLER:

Postman: GET /api/teachers/GV001
              ↓
@GetMapping("/{id}")
getTeacherById(@PathVariable String id)
              ↓
    id = "GV001"
              ↓
giangVienRepository.findById("GV001")
              ↓
    ┌─────────────────┐
    │  Tìm thấy?      │
    └────┬────────────┘
    YES  │         NO
    ↓    │         ↓
200 OK   │      404 Not Found
+ data   │      (không có body)
''')
code_block('''// Từ dự án thực tế - GiangVienController
@RestController
@RequestMapping("/teachers")
@RequiredArgsConstructor
@CrossOrigin(origins = "*")
public class GiangVienController {

    private final GiangVienRepository giangVienRepository;

    // GET /teachers → lấy tất cả
    @GetMapping
    public ResponseEntity<List<GiangVien>> getAllTeachers() {
        return ResponseEntity.ok(giangVienRepository.findAll());
    }

    // GET /teachers/GV001 → lấy theo mã
    @GetMapping("/{id}")
    public ResponseEntity<?> getTeacherById(@PathVariable String id) {
        return giangVienRepository.findById(id)
            .map(ResponseEntity::ok)           // tìm thấy → 200
            .orElse(ResponseEntity.notFound().build()); // không thấy → 404
    }

    // POST /teachers → thêm mới
    @PostMapping
    public ResponseEntity<?> createTeacher(@RequestBody GiangVien gv) {
        // check trùng mã
        if (giangVienRepository.existsById(gv.getMaGV())) {
            Map<String, Object> res = new HashMap<>();
            res.put("success", false);
            res.put("message", "Mã GV đã tồn tại");
            return ResponseEntity.badRequest().body(res); // 400
        }
        GiangVien saved = giangVienRepository.save(gv);
        Map<String, Object> res = new HashMap<>();
        res.put("success", true);
        res.put("data", saved);
        return ResponseEntity.ok(res); // 200
    }

    // DELETE /teachers/GV001 → xóa
    @DeleteMapping("/{id}")
    public ResponseEntity<?> deleteTeacher(@PathVariable String id) {
        return giangVienRepository.findById(id)
            .map(gv -> {
                giangVienRepository.delete(gv);
                Map<String, Object> res = new HashMap<>();
                res.put("success", true);
                return ResponseEntity.ok(res);
            })
            .orElse(ResponseEntity.notFound().build());
    }
}
''')

# 2.6 HTTP Status Codes
h2('2.6 HTTP Status Codes')
diagram('''
┌──────┬─────────────────────┬──────────────────────────────────┐
│ Code │ Tên                 │ Khi nào dùng                     │
├──────┼─────────────────────┼──────────────────────────────────┤
│ 200  │ OK                  │ Request thành công               │
│ 201  │ Created             │ Tạo mới thành công               │
├──────┼─────────────────────┼──────────────────────────────────┤
│ 400  │ Bad Request         │ Dữ liệu gửi lên sai/thiếu        │
│ 401  │ Unauthorized        │ Chưa đăng nhập / sai mật khẩu   │
│ 403  │ Forbidden           │ Không có quyền truy cập          │
│ 404  │ Not Found           │ Không tìm thấy tài nguyên        │
├──────┼─────────────────────┼──────────────────────────────────┤
│ 500  │ Internal Server Err │ Lỗi code phía server (bug)       │
└──────┴─────────────────────┴──────────────────────────────────┘

Ví dụ thực tế trong dự án:
  GET  /students/SV001        → 200 (tìm thấy)
  GET  /students/SV999        → 404 (không tồn tại)
  POST /students (trùng mã)   → 400 (mã đã tồn tại)
  POST /auth/login (sai pass) → 401 (sai mật khẩu)
''')

# 2.7 application.yml
h2('2.7 Cấu hình application.yml')
code_block('''# Từ dự án thực tế
server:
  port: 8080          # server chạy ở port 8080
  servlet:
    context-path: /api  # tất cả URL bắt đầu bằng /api

spring:
  datasource:
    url: jdbc:sqlserver://localhost:1433;databaseName=QuanLyDiemSinhVien
    username: sa
    password: yourpassword
    driver-class-name: com.microsoft.sqlserver.jdbc.SQLServerDriver

  jpa:
    hibernate:
      ddl-auto: none    # không tự tạo/xóa bảng
    show-sql: true      # in SQL ra console để debug
''')

# 2.8 Luồng đăng nhập
h2('2.8 Luồng Đăng Nhập Thực Tế (Dự án)')
diagram('''
POSTMAN / Desktop App
  POST /api/auth/login
  Body: {"username":"admin","password":"admin123"}
              ↓
  ┌───────────────────────────────────────────┐
  │  Spring Boot Filter Layer                 │
  │  Kiểm tra Content-Type, format request    │
  │  Nếu lỗi → 400 Bad Request (dừng ở đây)  │
  └──────────────────┬────────────────────────┘
                     ↓ (request hợp lệ)
  ┌───────────────────────────────────────────┐
  │  AuthController.login()                   │
  │  1. Validate username (regex)             │
  │     Nếu sai → 401 (dừng ở đây)           │
  │  2. Tìm user trong DB                     │
  └──────────────────┬────────────────────────┘
                     ↓
  ┌───────────────────────────────────────────┐
  │  TaiKhoanRepository.findById(username)    │
  │  SELECT * FROM TAIKHOAN WHERE             │
  │  ten_dang_nhap = 'admin'                  │
  └──────────────────┬────────────────────────┘
                     ↓
  ┌───────────────────────────────────────────┐
  │  SQL Server                               │
  │  Tìm thấy → trả TaiKhoan object          │
  │  Không thấy → trả empty                  │
  └──────────────────┬────────────────────────┘
                     ↓
  ┌───────────────────────────────────────────┐
  │  AuthController (tiếp)                    │
  │  So sánh password SHA-256                 │
  │  Đúng  → 200 + {success:true, role:Admin} │
  │  Sai   → 401 + {success:false}            │
  └──────────────────┬────────────────────────┘
                     ↓
  POSTMAN nhận response
  Script test: kiểm tra status + success
  → PASS hoặc FAIL
''')

# Save
doc.save('postman-tests/JavaCore_SpringBoot.docx')
print('Done - saved JavaCore_SpringBoot.docx')
