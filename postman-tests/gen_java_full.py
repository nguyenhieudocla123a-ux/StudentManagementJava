from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()
section = doc.sections[0]
section.left_margin = Cm(2)
section.right_margin = Cm(2)
section.top_margin = Cm(2)
section.bottom_margin = Cm(2)

def h1(text):
    p = doc.add_paragraph()
    run = p.add_run('═══ ' + text + ' ═══')
    run.font.name = 'Consolas'; run.font.size = Pt(15); run.bold = True
    run.font.color.rgb = RGBColor(0, 70, 127)

def h2(text):
    p = doc.add_paragraph()
    run = p.add_run('▶ ' + text)
    run.font.name = 'Consolas'; run.font.size = Pt(12); run.bold = True
    run.font.color.rgb = RGBColor(0, 112, 192)

def h3(text):
    p = doc.add_paragraph()
    run = p.add_run('  ◆ ' + text)
    run.font.name = 'Consolas'; run.font.size = Pt(11); run.bold = True
    run.font.color.rgb = RGBColor(68, 114, 196)

def txt(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Times New Roman'; run.font.size = Pt(11)

def code(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(text)
    run.font.name = 'Courier New'; run.font.size = Pt(9)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), 'F0F0F0')
    pPr.append(shd)

def box(text, color='E8F4FD'):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    run = p.add_run(text)
    run.font.name = 'Courier New'; run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0, 80, 0)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), color)
    pPr.append(shd)

def nb(): doc.add_paragraph()

# TITLE
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run('JAVA CORE → SPRING BOOT\nTÀI LIỆU HỌC ĐẦY ĐỦ')
r.font.name = 'Consolas'; r.font.size = Pt(20); r.bold = True
r.font.color.rgb = RGBColor(0, 70, 127)
nb()

# MINDMAP
h1('TỔNG QUAN LỘ TRÌNH')
box('''
                    ┌─────────────────┐
                    │   JAVA CORE     │
                    └────────┬────────┘
          ┌─────────┬────────┼────────┬─────────┐
          ↓         ↓        ↓        ↓         ↓
      ┌───────┐ ┌──────┐ ┌──────┐ ┌──────┐ ┌──────────┐
      │ Basic │ │ OOP  │ │Coll- │ │Excep-│ │ Java 8+  │
      │ Java  │ │      │ │ection│ │tion  │ │Lambda    │
      └───────┘ └──────┘ └──────┘ └──────┘ │Stream    │
                                            └──────────┘
                    ↓ (sau khi nắm Java Core)
                    ┌─────────────────┐
                    │  SPRING BOOT    │
                    └────────┬────────┘
          ┌─────────┬────────┼────────┬─────────┐
          ↓         ↓        ↓        ↓         ↓
      ┌───────┐ ┌──────┐ ┌──────┐ ┌──────┐ ┌──────┐
      │Entity │ │Repo- │ │Cont- │ │REST  │ │Config│
      │       │ │sitory│ │roller│ │API   │ │yml   │
      └───────┘ └──────┘ └──────┘ └──────┘ └──────┘
''', 'FFF9C4')
nb()

# ============================================================
doc.add_page_break()
h1('PHẦN 1: BASIC JAVA')

h2('1.1 Cấu trúc chương trình Java')
box('''
┌─────────────────────────────────────────────────────────┐
│  MyApp.java                                             │
│  ─────────────────────────────────────────────────────  │
│  package com.example;        ← tên gói (thư mục)       │
│                                                         │
│  import java.util.ArrayList; ← import thư viện         │
│                                                         │
│  public class MyApp {        ← tên class = tên file    │
│      public static void main(String[] args) {           │
│          System.out.println("Hello Java");              │
│      }                       ← điểm bắt đầu chạy      │
│  }                                                      │
└─────────────────────────────────────────────────────────┘
''')
txt('Compile và chạy:')
code('''javac MyApp.java   // compile → tạo MyApp.class
java MyApp         // chạy chương trình''')
nb()

h2('1.2 Kiểu dữ liệu')
box('''
PRIMITIVE (lưu giá trị trực tiếp)        REFERENCE (lưu địa chỉ)
─────────────────────────────────        ──────────────────────────
int     age = 25;                        String name = "Nguyen A";
long    bigNum = 9999999999L;            int[] arr = {1, 2, 3};
float   price = 9.99f;                   GiangVien gv = new GiangVien();
double  pi = 3.14159;                    List<String> list = new ArrayList<>();
char    grade = 'A';
boolean isPass = true;
byte    b = 127;
short   s = 32767;

Stack Memory                             Heap Memory
┌──────────┐                             ┌──────────────────┐
│ age = 25 │                             │ "Nguyen A"       │
│ pi = 3.14│                             │ GiangVien object │
│ name ────┼────────────────────────────►│   maGV="GV001"   │
└──────────┘                             │   hoTen="Nguyen" │
                                         └──────────────────┘
''')
nb()

h2('1.3 Biến và Hằng')
code('''// Biến thông thường - có thể thay đổi
int count = 0;
count = 5;  // OK

// Hằng số - KHÔNG thể thay đổi
final double PI = 3.14159;
// PI = 3.0;  // LỖI!

// Scope của biến
public class Example {
    static int staticVar = 10;    // biến static - dùng chung cho cả class
    int instanceVar = 20;         // biến instance - mỗi object có 1 bản

    void method() {
        int localVar = 30;        // biến local - chỉ tồn tại trong method
        System.out.println(localVar);
    }
}''')
nb()

h2('1.4 Toán tử')
box('''
┌──────────────┬────────────────────────────────────────────┐
│ Loại         │ Ví dụ                                      │
├──────────────┼────────────────────────────────────────────┤
│ Arithmetic   │ 5+3=8  5-3=2  5*3=15  10/3=3  10%3=1      │
│ Comparison   │ 5==5 → true   5!=3 → true   5>3 → true    │
│ Logical      │ true&&false→false  true||false→true  !true→false │
│ Assignment   │ x=5  x+=3(x=8)  x-=2(x=6)  x*=2(x=12)   │
│ Increment    │ x++ (dùng rồi tăng)  ++x (tăng rồi dùng) │
└──────────────┴────────────────────────────────────────────┘
''')
code('''int x = 5;
System.out.println(x++); // in 5, sau đó x=6
System.out.println(++x); // x=7, sau đó in 7

// Ternary operator
int max = (a > b) ? a : b;  // nếu a>b thì max=a, ngược lại max=b''')
nb()

h2('1.5 Cấu trúc điều khiển')
code('''// IF - ELSE
int diem = 8;
if (diem >= 8.5) {
    System.out.println("Xuất sắc");
} else if (diem >= 7.0) {
    System.out.println("Khá");
} else if (diem >= 5.0) {
    System.out.println("Trung bình");
} else {
    System.out.println("Yếu");
}

// SWITCH
String loai = "Admin";
switch (loai) {
    case "Admin":
        System.out.println("Quản trị viên");
        break;
    case "GiangVien":
        System.out.println("Giảng viên");
        break;
    default:
        System.out.println("Sinh viên");
}

// FOR loop
for (int i = 0; i < 5; i++) {
    System.out.println("Lần " + i);
}

// WHILE loop
int i = 0;
while (i < 5) {
    System.out.println(i);
    i++;
}

// FOR-EACH (dùng với Collection)
List<String> names = Arrays.asList("An", "Binh", "Cuong");
for (String name : names) {
    System.out.println(name);
}''')
nb()

# ============================================================
doc.add_page_break()
h1('PHẦN 2: LẬP TRÌNH HƯỚNG ĐỐI TƯỢNG (OOP)')

h2('2.1 Class và Object')
box('''
CLASS = Bản thiết kế (blueprint)
OBJECT = Thực thể được tạo ra từ class

┌─────────────────────────────────┐
│  class SinhVien {               │
│    // Attributes (thuộc tính)   │
│    String maSV;                 │
│    String hoTen;                │
│    float diemTB;                │
│                                 │
│    // Constructor               │
│    SinhVien(String maSV, ...) { │
│      this.maSV = maSV;          │
│    }                            │
│                                 │
│    // Methods (phương thức)     │
│    String getMaSV() {...}       │
│    void hoc() {...}             │
│  }                              │
└─────────────────────────────────┘
         ↓ new SinhVien(...)
┌──────────────┐  ┌──────────────┐  ┌──────────────┐
│  sv1         │  │  sv2         │  │  sv3         │
│  maSV=SV001  │  │  maSV=SV002  │  │  maSV=SV003  │
│  hoTen=An    │  │  hoTen=Binh  │  │  hoTen=Cuong │
└──────────────┘  └──────────────┘  └──────────────┘
''')
code('''public class SinhVien {
    // Attributes
    private String maSV;
    private String hoTen;
    private float diemTB;

    // Constructor không tham số
    public SinhVien() {}

    // Constructor có tham số
    public SinhVien(String maSV, String hoTen) {
        this.maSV = maSV;    // this = object hiện tại
        this.hoTen = hoTen;
    }

    // Getter
    public String getMaSV() { return maSV; }
    public String getHoTen() { return hoTen; }

    // Setter
    public void setMaSV(String maSV) { this.maSV = maSV; }

    // Method
    public String xepLoai() {
        if (diemTB >= 8.5) return "Xuất sắc";
        if (diemTB >= 7.0) return "Khá";
        return "Trung bình";
    }
}

// Tạo object
SinhVien sv1 = new SinhVien("SV001", "Nguyen Van An");
SinhVien sv2 = new SinhVien();
sv2.setMaSV("SV002");
System.out.println(sv1.getMaSV()); // SV001''')
nb()

h2('2.2 Encapsulation (Đóng gói)')
box('''
Che giấu dữ liệu bên trong, chỉ cho phép truy cập qua getter/setter

┌─────────────────────────────────────────────────────┐
│  class BankAccount {                                │
│    private double balance;  ← PRIVATE: không ai    │
│                               truy cập trực tiếp   │
│    public void deposit(double amount) {             │
│      if (amount > 0) balance += amount; ← kiểm tra │
│    }                                                │
│    public double getBalance() { return balance; }   │
│  }                                                  │
└─────────────────────────────────────────────────────┘

❌ SAI:  account.balance = -1000;  // không được!
✅ ĐÚNG: account.deposit(1000);    // qua method có kiểm tra
''')
nb()

h2('2.3 Inheritance (Kế thừa)')
box('''
class Animal (cha)              class Dog extends Animal (con)
┌──────────────────┐            ┌──────────────────────────────┐
│ String name      │            │ String breed                 │
│ int age          │  ←kế thừa─ │ (có cả name, age từ Animal) │
│ void eat() {...} │            │ void bark() {...}            │
│ void sleep(){...}│            │ @Override                    │
└──────────────────┘            │ void eat() { "ăn xương" }   │
                                └──────────────────────────────┘
''')
code('''// Class cha
public class NguoiDung {
    protected String tenDangNhap;
    protected String matKhau;

    public void dangNhap() {
        System.out.println("Đăng nhập: " + tenDangNhap);
    }
}

// Class con kế thừa
public class SinhVien extends NguoiDung {
    private String maSV;

    @Override  // ghi đè method của cha
    public void dangNhap() {
        super.dangNhap();  // gọi method của cha
        System.out.println("Chào sinh viên " + maSV);
    }
}

// Class con khác
public class GiangVien extends NguoiDung {
    private String maGV;
    // có thể dùng tenDangNhap, matKhau từ NguoiDung
}''')
nb()

h2('2.4 Polymorphism (Đa hình)')
box('''
OVERLOADING (cùng tên, khác tham số)    OVERRIDING (ghi đè method cha)
─────────────────────────────────────   ──────────────────────────────
void print(int x) {...}                 class Animal { void sound(){} }
void print(String s) {...}              class Dog extends Animal {
void print(int x, int y) {...}            @Override
                                          void sound() { "Gâu gâu" }
→ Java tự chọn đúng method              }
  dựa vào kiểu tham số                  class Cat extends Animal {
                                          @Override
                                          void sound() { "Meo meo" }
                                        }
''')
code('''// Polymorphism thực tế
Animal[] animals = {new Dog(), new Cat(), new Dog()};
for (Animal a : animals) {
    a.sound(); // tự gọi đúng method của từng loại
}
// Output: Gâu gâu, Meo meo, Gâu gâu''')
nb()

h2('2.5 Abstraction (Trừu tượng)')
code('''// Abstract class - không thể tạo object trực tiếp
public abstract class Shape {
    abstract double area();  // method trừu tượng - không có code

    void display() {         // method bình thường - có code
        System.out.println("Diện tích: " + area());
    }
}

// Interface - tất cả method đều trừu tượng
public interface Drawable {
    void draw();    // không có code
    void resize();  // không có code
}

// Class implement
public class Circle extends Shape implements Drawable {
    double radius;

    @Override
    double area() { return Math.PI * radius * radius; }

    @Override
    public void draw() { System.out.println("Vẽ hình tròn"); }

    @Override
    public void resize() { radius *= 2; }
}''')
nb()

# ============================================================
doc.add_page_break()
h1('PHẦN 3: STRING, ARRAY, ENUM')

h2('3.1 String')
box('''
String là IMMUTABLE (bất biến) - mỗi lần thay đổi tạo object mới

String s = "Hello";
s = s + " World";

Stack:  s ──────────────────────────────────────────────────►"Hello World"
Heap:   "Hello" (vẫn còn đó, chờ GC thu dọn)  "Hello World" (mới)
''')
code('''String s = "Hello Java";

// Các method thường dùng
s.length()           // 10 - độ dài
s.toUpperCase()      // "HELLO JAVA"
s.toLowerCase()      // "hello java"
s.trim()             // xóa khoảng trắng 2 đầu
s.contains("Java")   // true
s.startsWith("Hello")// true
s.endsWith("Java")   // true
s.replace("Java","Python") // "Hello Python"
s.split(" ")         // ["Hello", "Java"]
s.substring(0, 5)    // "Hello"
s.equals("Hello Java")     // true (so sánh nội dung)
s.equalsIgnoreCase("hello java") // true

// StringBuilder - MUTABLE, hiệu quả hơn khi nối nhiều chuỗi
StringBuilder sb = new StringBuilder();
sb.append("Hello");
sb.append(" ");
sb.append("Java");
String result = sb.toString(); // "Hello Java"

// String.format
String msg = String.format("Sinh viên %s có điểm %.2f", "SV001", 8.5);
// "Sinh viên SV001 có điểm 8.50"''')
nb()

h2('3.2 Array')
code('''// Array 1 chiều
int[] scores = new int[5];        // tạo mảng 5 phần tử
int[] scores2 = {8, 7, 9, 6, 8}; // khởi tạo luôn

scores[0] = 10;  // gán giá trị
System.out.println(scores[0]);    // đọc giá trị
System.out.println(scores.length); // 5

// Duyệt mảng
for (int i = 0; i < scores.length; i++) {
    System.out.println(scores[i]);
}
// Hoặc for-each
for (int score : scores) {
    System.out.println(score);
}

// Array 2 chiều (ma trận)
int[][] matrix = new int[3][3];
int[][] matrix2 = {{1,2,3}, {4,5,6}, {7,8,9}};
System.out.println(matrix2[1][2]); // 6 (hàng 1, cột 2)

// Arrays utility
Arrays.sort(scores);              // sắp xếp
Arrays.fill(scores, 0);           // điền giá trị
int[] copy = Arrays.copyOf(scores, 3); // copy 3 phần tử đầu''')
nb()

h2('3.3 Enum')
code('''// Định nghĩa enum
public enum LoaiNguoiDung {
    ADMIN,
    GIANG_VIEN,
    SINH_VIEN
}

// Dùng enum
LoaiNguoiDung loai = LoaiNguoiDung.ADMIN;

switch (loai) {
    case ADMIN:
        System.out.println("Quản trị viên");
        break;
    case GIANG_VIEN:
        System.out.println("Giảng viên");
        break;
    case SINH_VIEN:
        System.out.println("Sinh viên");
        break;
}

// Enum với thuộc tính
public enum DiemChu {
    A(8.5, "Xuất sắc"),
    B_PLUS(8.0, "Giỏi"),
    B(7.0, "Khá"),
    F(0.0, "Kém");

    private final double minScore;
    private final String moTa;

    DiemChu(double minScore, String moTa) {
        this.minScore = minScore;
        this.moTa = moTa;
    }
    public String getMoTa() { return moTa; }
}''')
nb()

# ============================================================
doc.add_page_break()
h1('PHẦN 4: EXCEPTION HANDLING')

h2('4.1 Luồng xử lý Exception')
box('''
LUỒNG BÌNH THƯỜNG:
code1 → code2 → code3 → code4 → kết thúc

LUỒNG KHI CÓ EXCEPTION:
code1 → code2 → LỖI! ──────────────────────────────────────►
                  ↓                                          │
              catch(Exception e)  ←─────────────────────────┘
                  ↓
              xử lý lỗi
                  ↓
              finally (luôn chạy dù có lỗi hay không)
                  ↓
              tiếp tục chương trình

CÁC LOẠI EXCEPTION:
┌─────────────────────────────────────────────────────────┐
│  Throwable                                              │
│  ├── Error (lỗi JVM - không xử lý được)                │
│  │   ├── OutOfMemoryError                              │
│  │   └── StackOverflowError                            │
│  └── Exception                                         │
│      ├── Checked (bắt buộc phải xử lý)                 │
│      │   ├── IOException                               │
│      │   └── SQLException                              │
│      └── Unchecked (RuntimeException)                  │
│          ├── NullPointerException                       │
│          ├── ArrayIndexOutOfBoundsException             │
│          ├── NumberFormatException                      │
│          └── IllegalArgumentException                   │
└─────────────────────────────────────────────────────────┘
''')
code('''// try-catch-finally
try {
    int result = 10 / 0;  // ArithmeticException
} catch (ArithmeticException e) {
    System.out.println("Lỗi: " + e.getMessage()); // / by zero
} catch (Exception e) {
    System.out.println("Lỗi khác: " + e.getMessage());
} finally {
    System.out.println("Luôn chạy dù có lỗi hay không");
}

// throw - ném exception
public void validate(String maSV) {
    if (maSV == null || maSV.isEmpty()) {
        throw new IllegalArgumentException("Mã SV không được rỗng");
    }
}

// throws - khai báo method có thể ném exception
public void readFile(String path) throws IOException {
    // code đọc file
}

// Custom Exception
public class DuplicateKeyException extends RuntimeException {
    public DuplicateKeyException(String message) {
        super(message);
    }
}

// Dùng custom exception
if (repository.existsById(maSV)) {
    throw new DuplicateKeyException("Mã SV " + maSV + " đã tồn tại");
}''')
nb()

# ============================================================
doc.add_page_break()
h1('PHẦN 5: COLLECTION FRAMEWORK')

h2('5.1 Tổng quan Collection')
box('''
                    Collection (interface)
                    ┌──────────┬──────────┐
                    ↓          ↓          ↓
                  List        Set        Queue
                    │          │
          ┌─────────┤    ┌─────┤
          ↓         ↓    ↓     ↓
      ArrayList LinkedList HashSet TreeSet
      (mảng động)(danh sách)(không trùng)(có thứ tự)

                    Map (interface) - riêng biệt
                    ┌──────────┬──────────┐
                    ↓          ↓          ↓
                HashMap    TreeMap   LinkedHashMap
               (nhanh nhất)(có thứ tự)(thứ tự insert)
''')
code('''// LIST - có thứ tự, cho phép trùng
List<String> list = new ArrayList<>();
list.add("GV001");
list.add("GV002");
list.add("GV001");  // cho phép trùng
System.out.println(list.size());    // 3
System.out.println(list.get(0));    // GV001
list.remove("GV001");               // xóa phần tử đầu tiên
list.contains("GV002");             // true

// SET - không có thứ tự, KHÔNG cho phép trùng
Set<String> set = new HashSet<>();
set.add("GV001");
set.add("GV002");
set.add("GV001");  // bị bỏ qua (trùng)
System.out.println(set.size());     // 2

// MAP - key-value, key không trùng
Map<String, GiangVien> map = new HashMap<>();
map.put("GV001", gv1);
map.put("GV002", gv2);
GiangVien found = map.get("GV001");  // tìm theo key
map.containsKey("GV001");            // true
map.size();                          // 2

// Duyệt Map
for (Map.Entry<String, GiangVien> entry : map.entrySet()) {
    System.out.println(entry.getKey() + " → " + entry.getValue().getHoTen());
}''')
nb()

h2('5.2 Generics')
box('''
Không có Generics:                  Có Generics:
List list = new ArrayList();        List<String> list = new ArrayList<>();
list.add("Hello");                  list.add("Hello");
list.add(123);  // OK nhưng nguy    // list.add(123); → LỖI compile time!
String s = (String) list.get(0);    String s = list.get(0); // không cần cast
String s2 = (String) list.get(1);  // an toàn hơn
// → ClassCastException lúc runtime!
''')
code('''// Generic method
public <T> void printArray(T[] arr) {
    for (T item : arr) {
        System.out.println(item);
    }
}

// Generic class
public class Pair<K, V> {
    private K key;
    private V value;

    public Pair(K key, V value) {
        this.key = key;
        this.value = value;
    }
}

// Dùng
Pair<String, Integer> pair = new Pair<>("GV001", 100);''')
nb()

# ============================================================
doc.add_page_break()
h1('PHẦN 6: JAVA 8+ (LAMBDA & STREAM)')

h2('6.1 Lambda Expression')
box('''
TRƯỚC Java 8 (Anonymous class):        JAVA 8+ (Lambda):
─────────────────────────────          ─────────────────
list.sort(new Comparator<String>() {   list.sort((a, b) -> a.compareTo(b));
    @Override
    public int compare(String a,       // Ngắn gọn hơn nhiều!
                       String b) {
        return a.compareTo(b);
    }
});

Cú pháp Lambda:
(tham số) -> { code }
(a, b)    -> a + b          // 1 dòng, không cần return
(a, b)    -> { return a+b; } // nhiều dòng
()        -> System.out.println("Hello")  // không tham số
''')
code('''// Ví dụ Lambda thực tế
List<GiangVien> list = dao.getAllGiangVien();

// Duyệt
list.forEach(gv -> System.out.println(gv.getHoTen()));

// Sắp xếp theo tên
list.sort((gv1, gv2) -> gv1.getHoTen().compareTo(gv2.getHoTen()));

// Lọc
list.removeIf(gv -> gv.getMaKhoa() == null);''')
nb()

h2('6.2 Stream API')
box('''
Stream = luồng xử lý dữ liệu theo chuỗi

Collection → stream() → filter() → map() → collect()
                ↓           ↓         ↓          ↓
            tạo stream   lọc data  biến đổi  thu thập kết quả

Ví dụ:
[GV001,GV002,GV003,GV004,GV005]
    → filter(khoa==CNTT) → [GV001,GV003,GV005]
    → map(getHoTen)      → ["Nguyen A","Tran B","Le C"]
    → collect(toList)    → List<String>
''')
code('''List<GiangVien> list = dao.getAllGiangVien();

// filter - lọc
List<GiangVien> cnttGV = list.stream()
    .filter(gv -> "CNTT".equals(gv.getMaKhoa()))
    .collect(Collectors.toList());

// map - biến đổi
List<String> names = list.stream()
    .map(GiangVien::getHoTen)  // method reference
    .collect(Collectors.toList());

// anyMatch - kiểm tra có phần tử nào thỏa không
boolean hasGV001 = list.stream()
    .anyMatch(gv -> "GV001".equals(gv.getMaGV()));

// findFirst - tìm phần tử đầu tiên
Optional<GiangVien> gv001 = list.stream()
    .filter(gv -> "GV001".equals(gv.getMaGV()))
    .findFirst();

// count - đếm
long count = list.stream()
    .filter(gv -> "CNTT".equals(gv.getMaKhoa()))
    .count();

// sorted - sắp xếp
List<GiangVien> sorted = list.stream()
    .sorted((a, b) -> a.getHoTen().compareTo(b.getHoTen()))
    .collect(Collectors.toList());''')
nb()

h2('6.3 Optional')
code('''// Optional - tránh NullPointerException
Optional<GiangVien> optional = repository.findById("GV001");

// Kiểm tra có giá trị không
if (optional.isPresent()) {
    GiangVien gv = optional.get();
    System.out.println(gv.getHoTen());
}

// Cách ngắn gọn hơn
optional.ifPresent(gv -> System.out.println(gv.getHoTen()));

// Lấy giá trị hoặc default
GiangVien gv = optional.orElse(new GiangVien());

// Ném exception nếu không có
GiangVien gv2 = optional.orElseThrow(
    () -> new RuntimeException("Không tìm thấy GV")
);''')
nb()

# ============================================================
doc.add_page_break()
h1('PHẦN 7: SPRING BOOT - KIẾN TRÚC')

h2('7.1 Kiến trúc 3 tầng')
box('''
┌─────────────────────────────────────────────────────────────────┐
│                    CLIENT                                        │
│  Postman / Desktop App / Web Browser                            │
│  Gửi: HTTP Request (GET/POST/PUT/DELETE) + JSON body            │
└──────────────────────────┬──────────────────────────────────────┘
                           │ HTTP
                           ↓
┌──────────────────────────▼──────────────────────────────────────┐
│                 PRESENTATION LAYER (Controller)                  │
│  @RestController                                                 │
│  - Nhận HTTP request                                            │
│  - Validate dữ liệu đầu vào                                     │
│  - Gọi Business Logic                                           │
│  - Trả HTTP response (JSON)                                     │
│                                                                  │
│  AuthController    SinhVienController    GiangVienController    │
└──────────────────────────┬──────────────────────────────────────┘
                           │ Java method call
                           ↓
┌──────────────────────────▼──────────────────────────────────────┐
│                 DATA ACCESS LAYER (Repository)                   │
│  JpaRepository                                                   │
│  - Truy vấn database                                            │
│  - CRUD operations                                              │
│  - Tự động tạo SQL                                              │
│                                                                  │
│  TaiKhoanRepository  SinhVienRepository  GiangVienRepository    │
└──────────────────────────┬──────────────────────────────────────┘
                           │ SQL
                           ↓
┌──────────────────────────▼──────────────────────────────────────┐
│                    DATABASE                                      │
│  SQL Server / PostgreSQL / MySQL                                 │
│  TAIKHOAN  SINHVIEN  GIANGVIEN  MONHOC  LOPHOCPHAN  DIEM       │
└─────────────────────────────────────────────────────────────────┘
''')
nb()

h2('7.2 Entity - Ánh xạ Database')
box('''
DATABASE TABLE ←──────────────────────────────► JAVA ENTITY
─────────────────                               ─────────────────────
SINHVIEN                                        @Entity
  ma_sv    VARCHAR(10) PK          ←→           @Id @Column("ma_sv")
  ho_ten   NVARCHAR(100)           ←→           @Column("ho_ten")
  ngay_sinh DATE                   ←→           @Column("ngay_sinh")
  gioi_tinh NVARCHAR(10)           ←→           @Column("gioi_tinh")
  email    VARCHAR(100)            ←→           @Column("email")
  ma_khoa  VARCHAR(10) FK          ←→           @Column("ma_khoa")
                                                @ManyToOne (quan hệ)
''')
code('''@Data                           // Lombok: tự tạo getter/setter/toString
@Entity
@Table(name = "SINHVIEN")
public class SinhVien {

    @Id
    @Column(name = "ma_sv")
    private String maSV;

    @Column(name = "ho_ten")
    private String hoTen;

    @Column(name = "ngay_sinh")
    private LocalDate ngaySinh;

    @Column(name = "gioi_tinh")
    private String gioiTinh;

    @Column(name = "email")
    private String email;

    @Column(name = "ma_khoa")
    private String maKhoa;

    // Quan hệ Many-to-One: nhiều SV thuộc 1 Khoa
    @ManyToOne
    @JoinColumn(name = "ma_khoa", insertable = false, updatable = false)
    private Khoa khoa;
}''')
nb()

h2('7.3 Repository - JPA')
code('''// Chỉ cần khai báo interface, Spring tự tạo implementation
public interface SinhVienRepository
    extends JpaRepository<SinhVien, String> {

    // Spring tự tạo SQL từ tên method:
    List<SinhVien> findByHoTenContaining(String keyword);
    // → SELECT * FROM SINHVIEN WHERE ho_ten LIKE '%keyword%'

    List<SinhVien> findByMaKhoa(String maKhoa);
    // → SELECT * FROM SINHVIEN WHERE ma_khoa = ?

    List<SinhVien> findByMaKhoaAndGioiTinh(String maKhoa, String gioiTinh);
    // → SELECT * WHERE ma_khoa=? AND gioi_tinh=?

    // Custom query với @Query
    @Query("SELECT s FROM SinhVien s WHERE s.diemTB >= :minDiem")
    List<SinhVien> findByDiemTBGreaterThan(@Param("minDiem") float minDiem);
}

// Dùng trong Controller:
@RequiredArgsConstructor  // Lombok: tự inject qua constructor
public class SinhVienController {
    private final SinhVienRepository sinhVienRepository;

    // Các method có sẵn từ JpaRepository:
    sinhVienRepository.findAll()           // lấy tất cả
    sinhVienRepository.findById("SV001")   // tìm theo ID
    sinhVienRepository.save(sv)            // thêm/cập nhật
    sinhVienRepository.delete(sv)          // xóa
    sinhVienRepository.existsById("SV001") // kiểm tra tồn tại
    sinhVienRepository.count()             // đếm
}''')
nb()

h2('7.4 Controller - REST API đầy đủ')
code('''@RestController
@RequestMapping("/students")
@RequiredArgsConstructor
@CrossOrigin(origins = "*")
public class SinhVienController {

    private final SinhVienRepository sinhVienRepository;

    // ─── GET ALL ───────────────────────────────────────────
    // GET /students
    @GetMapping
    public ResponseEntity<List<SinhVien>> getAllStudents() {
        List<SinhVien> list = sinhVienRepository.findAll();
        return ResponseEntity.ok(list);  // 200 + data
    }

    // ─── GET BY ID ─────────────────────────────────────────
    // GET /students/SV001
    @GetMapping("/{id}")
    public ResponseEntity<?> getStudentById(@PathVariable String id) {
        return sinhVienRepository.findById(id)
            .map(sv -> ResponseEntity.ok((Object)sv))  // 200
            .orElse(ResponseEntity.notFound().build()); // 404
    }

    // ─── SEARCH ────────────────────────────────────────────
    // GET /students/search?keyword=Nguyen
    @GetMapping("/search")
    public ResponseEntity<List<SinhVien>> search(@RequestParam String keyword) {
        return ResponseEntity.ok(
            sinhVienRepository.findByHoTenContaining(keyword)
        );
    }

    // ─── CREATE ────────────────────────────────────────────
    // POST /students
    @PostMapping
    public ResponseEntity<?> createStudent(@RequestBody SinhVien sv) {
        Map<String, Object> response = new HashMap<>();

        // Kiểm tra trùng mã
        if (sinhVienRepository.existsById(sv.getMaSV())) {
            response.put("success", false);
            response.put("message", "Mã SV đã tồn tại");
            return ResponseEntity.badRequest().body(response); // 400
        }

        try {
            SinhVien saved = sinhVienRepository.save(sv);
            response.put("success", true);
            response.put("message", "Thêm thành công");
            response.put("data", saved);
            return ResponseEntity.ok(response); // 200
        } catch (Exception e) {
            response.put("success", false);
            response.put("message", "Lỗi: " + e.getMessage());
            return ResponseEntity.badRequest().body(response); // 400
        }
    }

    // ─── UPDATE ────────────────────────────────────────────
    // PUT /students/SV001
    @PutMapping("/{id}")
    public ResponseEntity<?> updateStudent(
            @PathVariable String id,
            @RequestBody SinhVien sv) {

        return sinhVienRepository.findById(id)
            .map(existing -> {
                sv.setMaSV(id);
                SinhVien updated = sinhVienRepository.save(sv);
                Map<String, Object> res = new HashMap<>();
                res.put("success", true);
                res.put("data", updated);
                return ResponseEntity.ok(res);
            })
            .orElse(ResponseEntity.notFound().build()); // 404
    }

    // ─── DELETE ────────────────────────────────────────────
    // DELETE /students/SV001
    @DeleteMapping("/{id}")
    public ResponseEntity<?> deleteStudent(@PathVariable String id) {
        return sinhVienRepository.findById(id)
            .map(sv -> {
                sinhVienRepository.delete(sv);
                Map<String, Object> res = new HashMap<>();
                res.put("success", true);
                res.put("message", "Xóa thành công");
                return ResponseEntity.ok(res);
            })
            .orElse(ResponseEntity.notFound().build()); // 404
    }
}''')
nb()

h2('7.5 application.yml - Cấu hình')
code('''server:
  port: 8080                    # port chạy server
  servlet:
    context-path: /api          # prefix cho tất cả URL

spring:
  datasource:
    url: jdbc:sqlserver://localhost:1433;databaseName=QuanLyDiemSinhVien;encrypt=false
    username: sa
    password: yourpassword
    driver-class-name: com.microsoft.sqlserver.jdbc.SQLServerDriver

  jpa:
    hibernate:
      ddl-auto: none            # none: không tự tạo/xóa bảng
                                # create: tạo bảng khi start
                                # update: cập nhật bảng
                                # create-drop: tạo rồi xóa khi stop
    show-sql: true              # in SQL ra console
    properties:
      hibernate:
        format_sql: true        # format SQL cho dễ đọc

logging:
  level:
    root: INFO
    com.studentmanagement: DEBUG  # log chi tiết cho package của mình''')
nb()

h2('7.6 Luồng Request đầy đủ')
box('''
POSTMAN gửi: POST http://localhost:8080/api/students
Body: {"maSV":"SV101","hoTen":"Nguyen Van Test","maKhoa":"CNTT"}

BƯỚC 1: Spring Boot nhận HTTP Request
        ↓
BƯỚC 2: Filter Chain
        - CORS filter: kiểm tra origin
        - Content-Type filter: kiểm tra JSON hợp lệ
        - Nếu lỗi → 400 Bad Request (dừng)
        ↓
BƯỚC 3: DispatcherServlet
        - Tìm Controller phù hợp với /students
        - Tìm method phù hợp với POST
        ↓
BƯỚC 4: SinhVienController.createStudent()
        - @RequestBody: Spring tự convert JSON → SinhVien object
        - Kiểm tra existsById("SV101")
        - Nếu trùng → return 400
        ↓
BƯỚC 5: SinhVienRepository.save(sv)
        - JPA tạo SQL: INSERT INTO SINHVIEN VALUES (?,?,?,...)
        ↓
BƯỚC 6: SQL Server thực thi INSERT
        - Thành công → trả SinhVien object đã lưu
        - Lỗi → throw Exception
        ↓
BƯỚC 7: Controller nhận kết quả
        - Tạo response: {success:true, data:{...}}
        - return ResponseEntity.ok(response)
        ↓
BƯỚC 8: Spring Boot gửi HTTP Response
        Status: 200 OK
        Body: {"success":true,"message":"Thêm thành công","data":{...}}
        ↓
POSTMAN nhận response
Tests script chạy: kiểm tra status=200, success=true → PASS
''', 'FFF3E0')

doc.save('postman-tests/JavaCore_SpringBoot_Full.docx')
print('Done - saved JavaCore_SpringBoot_Full.docx')
